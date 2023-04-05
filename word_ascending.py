from docx import Document
from tkinter import Tk
from tkinter.filedialog import askopenfilename
import os

def select_file():
    win = Tk()  # GUI 실행하고
    win.withdraw()
    docx = askopenfilename(title="워드 파일을 선택해주세요. by 우혁쌤",
                             initialdir=os.getcwd(),
                             filetypes=[("워드파일", "*.docx")])
    win.quit()  # GUI 종료
    return docx

def separate_list(lst):
    eng_lst = []
    kor_lst = []
    for item in lst:
        if isinstance(item, str) and item[0].isalpha():
            if ord(item[0]) < 256: # 영어로 시작하는 경우
                eng_lst.append(item)
            else: # 한글로 시작하는 경우
                kor_lst.append(item)
    return eng_lst, kor_lst

def text_to_list(data):
    for x, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        data.append(text)
    return data

def list_to_docx(list, document):
    for i in range(0, len(list)):
        text=list[i]
        document.add_paragraph(text)
    return document

if __name__ == '__main__':
    path=select_file()
    document = Document(path)
    data = []
    text_to_list(data)

    # 전처리(중복제거)
    data=set(data)
    data=list(data)
    # ''제거
    data.remove('')

    eng_lst, kor_lst=separate_list(data)
    kor_lst.sort()
    eng_lst.sort()

    new_document = Document()

    list_to_docx(kor_lst, new_document)
    list_to_docx(eng_lst, new_document)

    new_document.save(path.replace(".docx", "(정렬).docx"))
